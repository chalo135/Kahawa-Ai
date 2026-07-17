"""
RAG (Retrieval-Augmented Generation) module for Kahawa Smart API.

Workflow:
  1. Put your rust-disease source documents (.docx or .pdf) in the
     'Knowledge base/' folder.
  2. Run: python build_rag_index.py   (once, and again whenever docs change)
  3. main.py loads the resulting index at startup and retrieves relevant
     chunks for every /api/chat request automatically.

Embeddings come from Ollama's local 'nomic-embed-text' model, so the index
build and question-embedding both need Ollama running (verified available).
"""
import os
import pickle
import re
import numpy as np
import requests as req_lib
from docx import Document
from pypdf import PdfReader

OLLAMA_EMBED_URL = os.getenv("OLLAMA_EMBED_URL", "http://localhost:11434/api/embeddings")
EMBED_MODEL      = os.getenv("EMBED_MODEL", "nomic-embed-text")

INDEX_PATH    = "rag_index.pkl"
# Real folder name on disk (note the space + capital K).
KNOWLEDGE_DIR = os.getenv("KNOWLEDGE_DIR", "Knowledge base")

CHUNK_SIZE    = 800   # characters per chunk
CHUNK_OVERLAP = 150   # overlap between consecutive chunks
TOP_K         = 3     # chunks retrieved per chat message

# Cosine similarity below this means nothing in the knowledge base is really
# about the question — the caller must NOT pretend it has grounding.
# Calibrated against this actual index (nomic-embed-text scores):
#   on-topic  "how does shade affect rust?"  -> 0.898
#   on-topic  "what is Hemileia vastatrix?"  -> 0.526
#   off-topic "what is the capital of France?" -> 0.442
#   off-topic "how do I bake a cake?"          -> 0.413
# 0.50 keeps real questions and drops clearly unrelated ones. Retrieval is
# deliberately a little permissive; the prompt is what stops the model from
# answering when the retrieved text does not actually cover the question.
MIN_RELEVANCE = float(os.getenv("RAG_MIN_RELEVANCE", "0.50"))


def _clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _chunk_text(text: str, source: str) -> list[dict]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end]
        if chunk.strip():
            chunks.append({"text": chunk, "source": source})
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _embed(text: str) -> np.ndarray:
    resp = req_lib.post(
        OLLAMA_EMBED_URL,
        json={"model": EMBED_MODEL, "prompt": text},
        timeout=60,
    )
    resp.raise_for_status()
    return np.array(resp.json()["embedding"], dtype=np.float32)


def _read_docx(path: str) -> str:
    """Extracts text from a .docx: paragraphs plus any table cells."""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_document(path: str) -> str:
    if path.lower().endswith(".docx"):
        return _read_docx(path)
    if path.lower().endswith(".pdf"):
        return _read_pdf(path)
    return ""


def build_index():
    """Reads every .docx/.pdf in the knowledge base folder, chunks it, embeds
    each chunk with Ollama, and saves the result to disk. Run this manually —
    it is NOT called on startup, since embedding takes a while and should only
    happen when your documents actually change."""
    if not os.path.isdir(KNOWLEDGE_DIR):
        raise FileNotFoundError(
            f"Create a '{KNOWLEDGE_DIR}' folder next to main.py and put your "
            "source document(s) (.docx or .pdf) in it first."
        )

    doc_files = [
        f for f in os.listdir(KNOWLEDGE_DIR)
        if f.lower().endswith((".docx", ".pdf")) and not f.startswith("~$")
    ]
    if not doc_files:
        raise FileNotFoundError(
            f"No .docx or .pdf files found in '{KNOWLEDGE_DIR}'."
        )

    all_chunks = []
    for filename in sorted(doc_files):
        path = os.path.join(KNOWLEDGE_DIR, filename)
        print(f"Reading {filename}...")
        full_text = _clean_text(_read_document(path))
        if not full_text:
            print("  -> no readable text, skipped")
            continue
        chunks = _chunk_text(full_text, source=filename)
        all_chunks.extend(chunks)
        print(f"  -> {len(chunks)} chunks")

    if not all_chunks:
        raise RuntimeError("No readable text extracted from the knowledge base.")

    print(f"Embedding {len(all_chunks)} chunks with '{EMBED_MODEL}'... this may take a few minutes.")
    embeddings = []
    for i, chunk in enumerate(all_chunks):
        embeddings.append(_embed(chunk["text"]))
        if (i + 1) % 10 == 0:
            print(f"  embedded {i + 1}/{len(all_chunks)}")

    embeddings = np.vstack(embeddings)

    with open(INDEX_PATH, "wb") as f:
        pickle.dump({"chunks": all_chunks, "embeddings": embeddings}, f)

    print(f"Done. Saved index with {len(all_chunks)} chunks to {INDEX_PATH}")


_index_cache = None


def load_index():
    """Loads the saved index into memory. Call once at backend startup."""
    global _index_cache
    if not os.path.exists(INDEX_PATH):
        print(f"[RAG] WARNING: {INDEX_PATH} not found. RAG is disabled until you run: python build_rag_index.py")
        _index_cache = None
        return
    with open(INDEX_PATH, "rb") as f:
        _index_cache = pickle.load(f)
    print(f"[RAG] Index loaded: {len(_index_cache['chunks'])} chunks available")


def retrieve(query: str, k: int = TOP_K) -> list[dict]:
    """Returns the top-k most relevant chunks for a query.
    Returns an empty list if no index has been built/loaded yet,
    so /chat still works fine even before you set up RAG."""
    if _index_cache is None:
        return []

    try:
        query_vec = _embed(query)
    except req_lib.RequestException:
        # Embedding model unreachable/not pulled — fail gracefully, chat still works
        return []

    chunk_vecs = _index_cache["embeddings"]
    query_norm = query_vec / (np.linalg.norm(query_vec) + 1e-8)
    chunk_norms = chunk_vecs / (np.linalg.norm(chunk_vecs, axis=1, keepdims=True) + 1e-8)
    scores = chunk_norms @ query_norm

    top_indices = np.argsort(scores)[::-1][:k]
    results = []
    for idx in top_indices:
        results.append({
            "text": _index_cache["chunks"][idx]["text"],
            "source": _index_cache["chunks"][idx]["source"],
            "score": float(scores[idx]),
        })
    return results


def retrieve_relevant(query: str, k: int = TOP_K) -> list[dict]:
    """Like retrieve(), but only returns chunks that clear MIN_RELEVANCE.

    An empty list is a meaningful signal: the knowledge base has nothing
    about this question, so the caller must not present an answer as if it
    were grounded in the documents.
    """
    return [hit for hit in retrieve(query, k) if hit["score"] >= MIN_RELEVANCE]


def index_is_loaded() -> bool:
    return _index_cache is not None